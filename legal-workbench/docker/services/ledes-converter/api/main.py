from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from typing import Annotated
import docx
import re
import os
import tempfile
import logging
from datetime import datetime
from collections import defaultdict
import time
import magic
from defusedxml import ElementTree

from .models import LedesData, ConversionResponse, HealthResponse, LineItem

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="LEDES Converter API",
    description="Convert DOCX invoices to LEDES 1998B format",
    version="1.0.0"
)

# CORS middleware - configured for production security
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost,http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["Content-Type", "Authorization"],
)

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILENAME_LENGTH = 255
MAX_DESCRIPTION_LENGTH = 500
MAX_LINE_ITEMS = 1000
ALLOWED_MIME_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream"  # Some systems send DOCX as this
]

# Simple rate limiting (production should use Redis)
rate_limit_storage = defaultdict(list)


def rate_limit_check(client_ip: str, max_requests: int = 10, window_seconds: int = 60) -> bool:
    """Simple in-memory rate limiting. Production should use Redis."""
    now = time.time()
    # Clean old entries
    rate_limit_storage[client_ip] = [
        ts for ts in rate_limit_storage[client_ip]
        if now - ts < window_seconds
    ]

    if len(rate_limit_storage[client_ip]) >= max_requests:
        return False

    rate_limit_storage[client_ip].append(now)
    return True


def validate_file_type(content: bytes, filename: str) -> bool:
    """Validate file type using magic bytes (MIME detection)."""
    try:
        mime = magic.from_buffer(content, mime=True)

        # Check MIME type
        if mime not in ALLOWED_MIME_TYPES:
            logger.warning(f"Invalid MIME type detected: {mime} for file extension .docx")
            return False

        # Additional check: DOCX files should start with PK (ZIP header)
        if not content.startswith(b'PK'):
            logger.warning("File does not have valid ZIP/DOCX header")
            return False

        return True
    except Exception as e:
        logger.error(f"File type validation error: {e}")
        return False


def sanitize_string(value: str, max_length: int = MAX_DESCRIPTION_LENGTH) -> str:
    """Sanitize and truncate string values."""
    if not value:
        return ""
    # Remove control characters and excessive whitespace
    cleaned = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', value)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned[:max_length]


def parse_currency(value_str: str | None) -> float:
    """Parse currency string to float, removing symbols and formatting."""
    if not value_str:
        return 0.0
    # Remove 'US $', '$', ',', ' ' and other non-numeric chars except '.'
    clean_val = re.sub(r'[^\d.]', '', value_str)
    try:
        return float(clean_val)
    except ValueError:
        return 0.0


def format_date_ledes(date_str: str) -> str:
    """Convert date string to LEDES format (YYYYMMDD)."""
    try:
        # Assuming format "Month DD, YYYY" e.g., "December 15, 2025"
        dt = datetime.strptime(date_str, "%B %d, %Y")
        return dt.strftime("%Y%m%d")
    except ValueError:
        return ""


def extract_ledes_data(text: str) -> dict:
    """Extract invoice data from text content with validation."""
    data = {
        "invoice_date": "",
        "invoice_number": "",
        "client_id": "SALESFORCE",  # Placeholder based on template analysis
        "matter_id": "LITIGATION-BRAZIL",  # Placeholder
        "invoice_total": 0.0,
        "line_items": []
    }

    # Regex patterns (case-insensitive for robustness)
    date_pattern = re.compile(r"Date\s*of\s*Issuance:\s*(.*)", re.IGNORECASE)
    invoice_num_pattern = re.compile(r"Invoice\s*#\s*(\d+)", re.IGNORECASE)
    total_pattern = re.compile(r"Total\s*Gross\s*Amount:\s*(?:US\s*)?\$?([\d,]+\.?\d*)", re.IGNORECASE)

    # Line items pattern
    line_item_pattern = re.compile(r"(.*?)\s*US\s*\$([\d,]+)(?:\s|$)")

    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Extract Header Info
        date_match = date_pattern.search(line)
        if date_match:
            raw_date = date_match.group(1).strip()
            data["invoice_date"] = format_date_ledes(sanitize_string(raw_date, 50))

        inv_num_match = invoice_num_pattern.search(line)
        if inv_num_match:
            data["invoice_number"] = sanitize_string(inv_num_match.group(1).strip(), 50)

        total_match = total_pattern.search(line)
        if total_match:
            data["invoice_total"] = parse_currency(total_match.group(1))

        # Extract Line Items (with limit to prevent DoS)
        if len(data["line_items"]) >= MAX_LINE_ITEMS:
            logger.warning(f"Reached max line items limit ({MAX_LINE_ITEMS})")
            break

        if "Total Gross Amount" not in line:
            item_match = line_item_pattern.search(line)
            if item_match:
                desc = sanitize_string(item_match.group(1).strip())
                amount = parse_currency(item_match.group(2))

                # Filter out likely false positives
                if desc and amount > 0:
                    data["line_items"].append({
                        "description": desc,
                        "amount": amount
                    })

    return data


def generate_ledes_1998b(data: dict) -> str:
    """Generate LEDES 1998B format from extracted data."""
    # LEDES 1998B Header
    header = (
        "INVOICE_DATE|INVOICE_NUMBER|CLIENT_ID|MATTER_ID|INVOICE_TOTAL|"
        "BILLING_START_DATE|BILLING_END_DATE|INVOICE_DESCRIPTION|"
        "LINE_ITEM_NUMBER|EXP/FEE/INV_ADJ_TYPE|LINE_ITEM_DATE|"
        "LINE_ITEM_TASK_CODE|LINE_ITEM_EXPENSE_CODE|TIMEKEEPER_ID|"
        "LINE_ITEM_DESCRIPTION|LINE_ITEM_UNITS|LINE_ITEM_RATE|"
        "LINE_ITEM_ADJUSTMENT_AMOUNT|LINE_ITEM_TOTAL"
    )

    lines = [header]

    for i, item in enumerate(data["line_items"], 1):
        row = [
            data["invoice_date"],           # 1
            data["invoice_number"],         # 2
            data["client_id"],              # 3
            data["matter_id"],              # 4
            f"{data['invoice_total']:.2f}", # 5
            "",                             # 6
            "",                             # 7
            "Legal Services",               # 8
            str(i),                         # 9
            "F",                            # 10
            "",                             # 11
            "",                             # 12
            "",                             # 13
            "",                             # 14
            item["description"],            # 15
            "",                             # 16
            "",                             # 17
            "",                             # 18
            f"{item['amount']:.2f}"         # 19
        ]

        lines.append("|".join(row))

    return "\n".join(lines)


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint for Docker/Kubernetes."""
    return HealthResponse(status="ok", service="ledes-converter")


@app.post("/convert/docx-to-ledes", response_model=ConversionResponse)
async def convert_docx_to_ledes(
    request: Request,
    file: Annotated[UploadFile, File(description="DOCX file to convert to LEDES")]
) -> ConversionResponse:
    """
    Convert a DOCX invoice file to LEDES 1998B format.

    - **file**: DOCX file containing invoice data (max 10MB)
    - Returns: LEDES formatted content and extracted data

    Security features:
    - Rate limiting (10 requests/minute per IP)
    - MIME type validation using magic bytes
    - File size validation
    - Input sanitization
    """
    # Rate limiting
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limit_check(client_ip):
        raise HTTPException(
            status_code=429,
            detail="Too many requests. Please try again later."
        )

    # Validate filename
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required.")

    if len(file.filename) > MAX_FILENAME_LENGTH:
        raise HTTPException(status_code=400, detail="Filename too long.")

    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .docx file.")

    # Read and validate file size
    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB."
        )

    # Validate MIME type using magic bytes
    if not validate_file_type(content, file.filename):
        raise HTTPException(
            status_code=400,
            detail="Invalid file format. File does not appear to be a valid DOCX document."
        )

    tmp_path = ""
    try:
        # Save temp file with secure permissions
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode='wb') as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # Set restrictive permissions
        os.chmod(tmp_path, 0o600)

        # Process DOCX with error handling
        try:
            doc = docx.Document(tmp_path)
        except Exception as docx_error:
            logger.error(f"Failed to parse DOCX: {docx_error}")
            raise HTTPException(
                status_code=400,
                detail="Unable to parse DOCX file. File may be corrupted or in an unsupported format."
            )

        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)

        text_content = "\n".join(full_text)

        # Validate extracted content
        if not text_content.strip():
            raise HTTPException(
                status_code=400,
                detail="No text content found in DOCX file."
            )

        # Extract Data
        extracted_data = extract_ledes_data(text_content)

        # Validate extracted data
        if not extracted_data.get("invoice_number"):
            logger.warning("Invoice number not found in document")

        if not extracted_data.get("line_items"):
            raise HTTPException(
                status_code=400,
                detail="No line items found in invoice. Please check document format."
            )

        # Generate LEDES
        ledes_content = generate_ledes_1998b(extracted_data)

        logger.info(f"Successfully converted invoice (lines: {len(extracted_data['line_items'])})")

        # Create validated response using Pydantic models
        line_items = [LineItem(**item) for item in extracted_data["line_items"]]
        ledes_data = LedesData(
            invoice_date=extracted_data["invoice_date"],
            invoice_number=extracted_data["invoice_number"],
            client_id=extracted_data["client_id"],
            matter_id=extracted_data["matter_id"],
            invoice_total=extracted_data["invoice_total"],
            line_items=line_items
        )

        return ConversionResponse(
            filename=sanitize_string(file.filename, MAX_FILENAME_LENGTH),
            status="success",
            extracted_data=ledes_data,
            ledes_content=ledes_content
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Conversion failed: {type(e).__name__}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Conversion failed due to an internal error. Please contact support if the issue persists."
        )
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception as cleanup_error:
                logger.error(f"Failed to cleanup temp file: {cleanup_error}")
